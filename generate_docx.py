import os
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches

def create_docx():
    doc = Document()
    doc.add_heading('College Electrical Engineering', 0)
    
    chapters = [
        ("college-electrical-engineering/ch01.html", "Chapter 1: Circuits Principles, Experiments, and Semiconductors"),
        ("college-electrical-engineering/ch02.html", "Chapter 2: Communication Circuits"),
        ("college-electrical-engineering/ch03.html", "Chapter 3: Antennae"),
        ("college-electrical-engineering/ch04.html", "Chapter 4: Digital Electronics")
    ]
    
    for html_file, title in chapters:
        if not os.path.exists(html_file):
            print(f"Skipping {html_file}, not found.")
            continue
            
        doc.add_heading(title, level=1)
        
        with open(html_file, 'r', encoding='utf-8') as f:
            soup = BeautifulSoup(f.read(), 'html.parser')
            
        # The content is within <div class="page-content">
        page_contents = soup.find_all('div', class_='page-content')
        
        for page in page_contents:
            for element in page.children:
                if element.name == 'p':
                    text = element.get_text(strip=True)
                    if text:
                        doc.add_paragraph(text)
                elif element.name == 'img':
                    src = element.get('src')
                    if src:
                        # src is like "images/crop_ch01_page_001_1.png"
                        img_path = os.path.join(os.path.dirname(html_file), src)
                        if os.path.exists(img_path):
                            try:
                                doc.add_picture(img_path, width=Inches(5.0))
                            except Exception as e:
                                print(f"Error adding image {img_path}: {e}")
                        else:
                            print(f"Image not found: {img_path}")
            
            # Add a page break after each page for clarity (optional, but good for manuscripts)
            # doc.add_page_break()

    output_path = "college-electrical-engineering/college_electrical_engineering_manuscript.docx"
    doc.save(output_path)
    print(f"Successfully generated {output_path}")

if __name__ == "__main__":
    create_docx()
